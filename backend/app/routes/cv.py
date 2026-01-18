from fastapi import APIRouter, UploadFile, File, HTTPException, Body  # FastAPI-komponenter för routes och fel
from fastapi.responses import JSONResponse, FileResponse  # JSON- och fil-respons
from typing import Optional  # Typ för valfria strängar
import tempfile  # Skapar temporära filer
import os  # Filhantering på OS-nivå
from pathlib import Path  # Path-verktyg för sökvägar
from docx import Document  # Skapar DOCX-filer från text
from app.services import pdf_exporter  # Skapar PDF-filer från text

from app.services import cv_reader, cv_storage  # Importerar CV-läsare och lagring
from app.models import JobPostingInput, KeywordRequest, OptimizeRequest  # Importerar request-modeller
from app.services import job_scraper, keyword_extractor, lm_studio_client  # Importerar webbskrapare, keywords och AI-klient

router = APIRouter(prefix="/api/cv", tags=["CV"])

def _normalize_input_text(value: Optional[str]) -> Optional[str]:  # Normaliserar input från Swagger eller klient
    if value is None:  # Om värdet saknas
        return None  # Returnerar None så vi kan falla tillbaka på sparad text
    stripped = value.strip()  # Tar bort onödiga blanksteg
    if stripped == "":  # Om strängen är tom
        return None  # Returnerar None för att signalera att vi ska använda sparad text
    if stripped.lower() == "string":  # Swagger placeholder-värde som inte är riktig input
        return None  # Returnerar None för att undvika felaktig prompt
    return stripped  # Returnerar normaliserad text

def _is_cookie_banner(text: str) -> bool:  # Enkel kontroll för cookie-banner-text
    lowered = text.lower()  # Gör texten gemener för enkel matchning
    return "kakor" in lowered or "cookies" in lowered  # Returnerar True om texten ser ut som cookie-banner

@router.post("/upload")
async def upload_cv(file: UploadFile = File()):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Ingen fil vald")
    
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]
    
    if file.content_type and file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Filtypen {file.content_type} stöds inte. Använd PDF eller DOCX."
        )
    
    file_suffix = Path(file.filename).suffix if file.filename else ""
    temp_file_path = None
    
    try:
        content = await file.read()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix) as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        cv_text = await cv_reader.read_cv_from_file(temp_file_path)
        
        saved_path = await cv_storage.save_master_cv(cv_text)
        
        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "text": cv_text,
            "length": len(cv_text),
            "saved_to": saved_path
        })
        
    except cv_reader.CVReadError as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ett fel uppstod: {str(e)}")
        
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except PermissionError:
                pass

@router.post("/job-posting")
async def submit_job_posting(input_data: JobPostingInput):
    if not input_data.text and not input_data.url:
        raise HTTPException(status_code=400, detail="Antingen text eller url måste anges")
    
    if input_data.text and input_data.url:
        raise HTTPException(status_code=400, detail="Ange endast text ELLER url, inte båda")
    
    try:
        if input_data.url:
            job_text = await job_scraper.scrape_job_posting(input_data.url)
        else:
            job_text = input_data.text
        
        saved_path = await cv_storage.save_job_posting(job_text)
        
        return JSONResponse(content={
            "success": True,
            "text": job_text,
            "length": len(job_text),
            "saved_to": saved_path
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/job-posting-raw")  # Tar emot raw text för jobbannons
async def submit_job_posting_raw(raw_text: str = Body(..., media_type="text/plain")):
    cleaned = raw_text.strip()  # Tar bort onödiga blanksteg
    if not cleaned:  # Säkerställer att texten inte är tom
        raise HTTPException(status_code=400, detail="Jobbannonsen är tom")  # Returnerar fel om tom
    if _is_cookie_banner(cleaned):  # Stoppar cookie-text
        raise HTTPException(status_code=400, detail="Jobbannonsen ser ut att vara cookie-text. Klistra in annonsen som text.")  # Tydligt fel
    saved_path = await cv_storage.save_job_posting(cleaned)  # Sparar jobbannonsen
    return JSONResponse(content={  # Returnerar resultat
        "success": True,  # Indikerar att sparning lyckades
        "length": len(cleaned),  # Längden på texten
        "saved_to": saved_path,  # Var filen sparades
    })  # Slut på svar

@router.post("/master-cv-raw")  # Tar emot raw text för CV
async def submit_master_cv_raw(raw_text: str = Body(..., media_type="text/plain")):
    cleaned = raw_text.strip()  # Tar bort onödiga blanksteg
    if not cleaned:  # Säkerställer att texten inte är tom
        raise HTTPException(status_code=400, detail="CV-texten är tom")  # Returnerar fel om tom
    saved_path = await cv_storage.save_master_cv(cleaned)  # Sparar CV-texten
    return JSONResponse(content={  # Returnerar resultat
        "success": True,  # Indikerar att sparning lyckades
        "length": len(cleaned),  # Längden på texten
        "saved_to": saved_path,  # Var filen sparades
    })  # Slut på svar

@router.get("/master-cv")
async def get_master_cv():
    cv_text = cv_storage.get_master_cv()
    if not cv_text:
        raise HTTPException(status_code=404, detail="Inget master CV hittades")
    return JSONResponse(content={
        "text": cv_text,
        "length": len(cv_text)
    })

@router.get("/job-posting")
async def get_job_posting():
    job_text = cv_storage.get_job_posting()
    if not job_text:
        raise HTTPException(status_code=404, detail="Ingen jobbannons hittades")
    return JSONResponse(content={
        "text": job_text,
        "length": len(job_text)
    })

@router.post("/keywords")  # Endpoint för att jämföra keywords
async def extract_keywords(input_data: KeywordRequest):  # Tar emot job_text/cv_text/top_n
    job_text = input_data.job_text or cv_storage.get_job_posting()  # Använder sparad jobbannons om ingen text skickas
    cv_text = input_data.cv_text or cv_storage.get_master_cv()  # Använder sparat CV om ingen text skickas
    if not job_text or not cv_text:  # Validerar att båda texterna finns
        raise HTTPException(status_code=400, detail="Saknar jobbannons eller CV")  # Returnerar fel vid saknad data
    result = keyword_extractor.compare_keywords(job_text, cv_text, input_data.top_n)  # Kör keyword-jämförelse
    return JSONResponse(content=result)  # Returnerar resultatet som JSON

@router.post("/optimize")  # Endpoint för AI-optimering av CV
async def optimize_cv(input_data: OptimizeRequest):  # Tar emot input för optimering
    job_text = _normalize_input_text(input_data.job_text) or cv_storage.get_job_posting()  # Hämtar jobbannons om den saknas i request
    cv_text = _normalize_input_text(input_data.cv_text) or cv_storage.get_master_cv()  # Hämtar CV om den saknas i request
    if not job_text or not cv_text:  # Validerar att båda texterna finns
        raise HTTPException(status_code=400, detail="Saknar jobbannons eller CV")  # Returnerar fel vid saknad data
    if _is_cookie_banner(job_text):  # Stoppar om jobbannonsen är cookie-banner
        raise HTTPException(status_code=400, detail="Jobbannonsen ser ut att vara cookie-text. Klistra in annonsen som text.")  # Tydligt felmeddelande
    optimized_text = await lm_studio_client.optimize_cv(cv_text, job_text)  # Anropar LM Studio för optimering
    return JSONResponse(content={  # Returnerar optimerat CV
        "optimized_cv": optimized_text,  # Den optimerade CV-texten
        "length": len(optimized_text),  # Antal tecken i svaret
    })  # Slut på JSON-svar

@router.post("/optimize-docx")  # Endpoint som returnerar optimerat CV som DOCX
async def optimize_cv_docx(input_data: OptimizeRequest):  # Tar emot input för optimering
    job_text = _normalize_input_text(input_data.job_text) or cv_storage.get_job_posting()  # Hämtar jobbannons om den saknas i request
    cv_text = _normalize_input_text(input_data.cv_text) or cv_storage.get_master_cv()  # Hämtar CV om den saknas i request
    if not job_text or not cv_text:  # Validerar att båda texterna finns
        raise HTTPException(status_code=400, detail="Saknar jobbannons eller CV")  # Returnerar fel vid saknad data
    if _is_cookie_banner(job_text):  # Stoppar om jobbannonsen är cookie-banner
        raise HTTPException(status_code=400, detail="Jobbannonsen ser ut att vara cookie-text. Klistra in annonsen som text.")  # Tydligt felmeddelande
    optimized_text = await lm_studio_client.optimize_cv(cv_text, job_text)  # Anropar LM Studio för optimering
    doc = Document()  # Skapar ett nytt DOCX-dokument
    lines = [line.strip() for line in optimized_text.splitlines()]  # Delar upp texten rad för rad
    for line in lines:  # Loopar igenom varje rad
        if not line:  # Hoppar över tomma rader
            continue  # Går vidare till nästa rad
        if line.startswith("=== ") and line.endswith(" ==="):  # Identifierar språk-sektioner
            doc.add_paragraph(line.replace("=", "").strip(), style="Heading 1")  # Lägger till huvudrubrik
        elif line.startswith("- ") or line.startswith("* "):  # Identifierar punktlistor
            doc.add_paragraph(line[2:], style="List Bullet")  # Lägger till bullet-paragraf
        elif line.isupper() or line.endswith(":"):  # Enkel heuristik för rubriker
            doc.add_paragraph(line, style="Heading 2")  # Lägger till rubrik
        else:  # Standardrad
            doc.add_paragraph(line)  # Lägger till vanlig paragraf
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")  # Skapar temporär fil
    temp_path = temp_file.name  # Sparar sökvägen till temp-filen
    temp_file.close()  # Stänger filen innan vi skriver DOCX (Windows-lås)
    doc.save(temp_path)  # Sparar DOCX till temp-filen
    return FileResponse(  # Returnerar DOCX-filen som download
        temp_path,  # Sökväg till filen
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # MIME-typ för DOCX
        filename="optimized_cv.docx",  # Filnamn för nedladdning
    )  # Slut på fil-respons

@router.post("/optimize-pdf")  # Endpoint som returnerar optimerat CV som PDF
async def optimize_cv_pdf(input_data: OptimizeRequest):  # Tar emot input för optimering
    job_text = _normalize_input_text(input_data.job_text) or cv_storage.get_job_posting()  # Hämtar jobbannons om den saknas i request
    cv_text = _normalize_input_text(input_data.cv_text) or cv_storage.get_master_cv()  # Hämtar CV om den saknas i request
    if not job_text or not cv_text:  # Validerar att båda texterna finns
        raise HTTPException(status_code=400, detail="Saknar jobbannons eller CV")  # Returnerar fel vid saknad data
    if _is_cookie_banner(job_text):  # Stoppar om jobbannonsen är cookie-banner
        raise HTTPException(status_code=400, detail="Jobbannonsen ser ut att vara cookie-text. Klistra in annonsen som text.")  # Tydligt felmeddelande
    optimized_text = await lm_studio_client.optimize_cv(cv_text, job_text)  # Anropar LM Studio för optimering
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")  # Skapar temporär PDF-fil
    temp_path = temp_file.name  # Hämtar temp-sökväg
    temp_file.close()  # Stänger filen innan vi skriver (Windows-lås)
    pdf_exporter.create_pdf_from_text(optimized_text, temp_path)  # Skapar PDF från text
    return FileResponse(  # Returnerar PDF-filen som download
        temp_path,  # Sökväg till filen
        media_type="application/pdf",  # MIME-typ för PDF
        filename="optimized_cv.pdf",  # Filnamn för nedladdning
    )  # Slut på fil-respons

